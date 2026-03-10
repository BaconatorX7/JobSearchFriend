import sys
from pathlib import Path
sys.path.append(str(Path(__file__).resolve().parents[2]))

import anthropic
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from config import ANTHROPIC_API_KEY, CLAUDE_MODEL, MAX_TOKENS, OUTPUT_DIR


client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)


# ─── Helpers ──────────────────────────────────────
def _ensure_output_dir() -> Path:
    """Create output directory if it doesn't exist."""
    path = Path(OUTPUT_DIR)
    path.mkdir(exist_ok=True)
    return path


def _safe_filename(company: str, title: str, prefix: str) -> str:
    """Generate a clean filename from company and title."""
    raw = f"{prefix}_{company}_{title}"
    return raw.replace(" ", "_").replace("/", "-")[:60] + ".docx"


def _save_docx(content: str, filepath: Path, title: str):
    """Save text content to a formatted .docx file."""
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1)
        section.right_margin  = Inches(1)

    # Title
    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.runs[0].font.size = Pt(14)

    doc.add_paragraph()  # spacing

    # Body — split on newlines to preserve paragraphs
    for line in content.split("\n"):
        para = doc.add_paragraph(line)
        para.runs[0].font.size = Pt(11) if para.runs else None

    doc.save(str(filepath))


# ─── Cover Letter ─────────────────────────────────
def generate_cover_letter(job: dict, resume_text: str = "") -> str:
    """Call Claude to generate a tailored cover letter."""
    prompt = f"""
You are an expert career coach and professional writer.

Write a compelling, concise cover letter (3 paragraphs, ~250 words) for this job.

## Job Details
- Title:    {job['title']}
- Company:  {job['company']}
- Location: {job['location']}
- URL:      {job['url']}

## Candidate Resume
{resume_text or "No resume provided — write a strong general cover letter."}

## Instructions
- Open with a strong hook, not "I am writing to apply..."
- Mirror keywords from the job title and industry
- Highlight 2-3 specific achievements or skills relevant to this role
- Close with a confident call to action
- Tone: professional but human
- Do NOT include placeholders like [Your Name]

Output only the cover letter text, ready to send.
"""
    message = client.messages.create(
        model=CLAUDE_MODEL,
        max_tokens=MAX_TOKENS,
        messages=[{"role": "user", "content": prompt}]
    )
    return message.content[0].text


# ─── Resume Tailor ────────────────────────────────
def tailor_resume(job: dict, resume_text: str = "") -> str:
    """Call Claude to tailor a resume for a specific job."""
    prompt = f"""
You are an expert resume writer and ATS optimization specialist.

Tailor this resume for the specific job below.

## Target Job
- Title:   {job['title']}
- Company: {job['company']}

## Original Resume
{resume_text or "No resume provided — generate a strong template resume."}

## Instructions
- Reorder and reword bullet points to match job keywords
- Quantify achievements where possible
- Optimize for ATS keyword matching
- Keep formatting clean and plain text friendly

Output the full tailored resume text, ready to copy-paste.
"""
    message = client.messages.create(
        model=CLAUDE_MODEL,
        max_tokens=MAX_TOKENS,
        messages=[{"role": "user", "content": prompt}]
    )
    return message.content[0].text


# ─── Main Entry ───────────────────────────────────
def generate_documents(job: dict, resume_text: str = "") -> tuple[str, str]:
    """
    Generate and save cover letter + tailored resume as .docx files.
    Returns (cover_letter_path, resume_path).
    """
    output_dir = _ensure_output_dir()

    # Cover letter
    cl_text     = generate_cover_letter(job, resume_text)
    cl_filename = _safe_filename(job["company"], job["title"], "cover_letter")
    cl_path     = output_dir / cl_filename
    _save_docx(cl_text, cl_path, f"Cover Letter — {job['title']} at {job['company']}")

    # Tailored resume
    res_text     = tailor_resume(job, resume_text)
    res_filename = _safe_filename(job["company"], job["title"], "resume")
    res_path     = output_dir / res_filename
    _save_docx(res_text, res_path, f"Resume — {job['title']} at {job['company']}")

    return str(cl_path), str(res_path)

# Test
if __name__ == "__main__":
    test_job = {
        "company": "Stripe",
        "title": "Backend Engineer",
        "location": "Remote",
        "url": "https://stripe.com/jobs"
    }
    cl, res = generate_documents(test_job, resume_text="")
    print(f"Cover letter: {cl}")
    print(f"Resume: {res}")